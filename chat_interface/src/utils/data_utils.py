import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def generate_conversation_id() -> str:
    """Generates a conversation ID with a UUID."""

    conversation_id = str(uuid.uuid4())
    return conversation_id


def chat_to_word(messages: dict) -> bytes:
    """Converts a chat to a Word document and returns the document as a byte string."""

    document = Document()

    style = document.styles["Normal"]
    font = style.font
    font.name = "Arial"

    section = document.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    image_path = Path(__file__).parent.parent / "images" / "neuralmind.png"
    run.add_picture(str(image_path), width=Inches(0.5))
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_heading("Chatbot Conversation History", 0)
    document.add_heading(
        f"Date and Time of Conversation: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
        level=1,
    )
    document.add_heading("Messages:\n", level=1)

    for message in messages:
        p = document.add_paragraph()
        if message["role"] == "assistant":
            role = p.add_run("Assistant: ")
            role.bold = True
            p.add_run(f"{message['content']}")
        elif message["role"] == "user":
            role = p.add_run("User: ")
            role.bold = True
            p.add_run(f"{message['content']}")

    word_file_io = BytesIO()
    document.save(word_file_io)
    word_file_io.seek(0)
    return word_file_io
